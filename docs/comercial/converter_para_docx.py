#!/usr/bin/env python3
"""Converte as propostas em markdown para .docx formatado.

Existe porque proposta comercial circula em Word: o cliente comenta, o jurídico
marca alteração, e alguém imprime. Markdown é ótimo para escrever e editar em
git — e inútil na mão de quem vai revisar a cláusula.

O conversor é deliberadamente pequeno e cobre só o que estes documentos usam:
títulos, parágrafos, negrito, itálico, código, tabelas, listas, citações e
regras horizontais. Não é um renderizador de markdown de propósito geral, e não
deve virar um: no dia em que a proposta precisar de algo que ele não faz, é mais
barato acrescentar aqui do que trocar por uma dependência que traz mil coisas.

Uso:
    python docs/comercial/converter_para_docx.py                 # converte todos
    python docs/comercial/converter_para_docx.py arquivo.md ...  # só estes
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PASTA = Path(__file__).resolve().parent

# A mesma paleta da apresentação: tinta com viés azul e a ressalva em ocre.
TINTA = RGBColor(0x10, 0x16, 0x20)
TINTA_FRACA = RGBColor(0x4C, 0x57, 0x68)
RESSALVA = RGBColor(0xA2, 0x57, 0x0C)

CORPO = "Calibri"
TITULO = "Calibri"


def _sombrear(celula, cor_hex: str) -> None:
    """Preenche o fundo de uma célula. python-docx não expõe isso direto."""
    elemento = OxmlElement("w:shd")
    elemento.set(qn("w:val"), "clear")
    elemento.set(qn("w:fill"), cor_hex)
    celula._tc.get_or_add_tcPr().append(elemento)


# Negrito, itálico e código, nesta ordem. O código vem por último para que
# `**texto**` dentro de crase não seja interpretado como negrito.
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")


def _escrever_inline(paragrafo, texto: str, tamanho: float | None = None) -> None:
    """Aplica negrito, itálico e código dentro de um parágrafo."""
    for parte in INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            r = paragrafo.add_run(parte[2:-2])
            r.bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            r = paragrafo.add_run(parte[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RESSALVA
        elif parte.startswith("*") and parte.endswith("*"):
            r = paragrafo.add_run(parte[1:-1])
            r.italic = True
        else:
            r = paragrafo.add_run(parte)
        r.font.name = r.font.name or CORPO
        if tamanho:
            r.font.size = Pt(tamanho)


def _titulo(doc, texto: str, nivel: int):
    tamanhos = {1: 20, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if nivel <= 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", texto))
    r.bold = True
    r.font.name = TITULO
    r.font.size = Pt(tamanhos.get(nivel, 11))
    r.font.color.rgb = TINTA if nivel <= 2 else RESSALVA
    return p


def _linha_de_tabela(linha: str) -> list[str]:
    return [c.strip() for c in linha.strip().strip("|").split("|")]


def _tabela(doc, linhas: list[str]) -> None:
    """Uma tabela markdown, com cabeçalho sombreado e alinhamento respeitado."""
    cabecalho = _linha_de_tabela(linhas[0])
    alinhamento = _linha_de_tabela(linhas[1])
    corpo = [_linha_de_tabela(l) for l in linhas[2:] if l.strip()]

    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    for indice, texto in enumerate(cabecalho):
        celula = tabela.rows[0].cells[indice]
        celula.text = ""
        _sombrear(celula, "EDEFF2")
        p = celula.paragraphs[0]
        _escrever_inline(p, texto, 9.5)
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = TINTA_FRACA

    for valores in corpo:
        celulas = tabela.add_row().cells
        for indice, texto in enumerate(valores[: len(cabecalho)]):
            celula = celulas[indice]
            celula.text = ""
            p = celula.paragraphs[0]
            # `---:` no separador significa coluna numérica, alinhada à direita.
            if indice < len(alinhamento) and alinhamento[indice].endswith(":"):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _escrever_inline(p, texto.replace("<br>", " "), 10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _citacao(doc, linhas: list[str]) -> None:
    """Bloco de instrução ou destaque, recuado e em ocre."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    _escrever_inline(p, " ".join(linhas), 10)
    for r in p.runs:
        r.font.color.rgb = RESSALVA


def converter(origem: Path) -> Path:
    texto = origem.read_text(encoding="utf-8")
    doc = Document()

    secao = doc.sections[0]
    secao.page_width, secao.page_height = Cm(21), Cm(29.7)
    for lado in ("left_margin", "right_margin"):
        setattr(secao, lado, Cm(2.4))
    secao.top_margin, secao.bottom_margin = Cm(2.2), Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = CORPO
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TINTA
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    linhas = texto.splitlines()
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        despida = linha.strip()

        if not despida:
            i += 1
            continue

        if despida.startswith("#"):
            nivel = len(despida) - len(despida.lstrip("#"))
            _titulo(doc, despida.lstrip("#").strip(), nivel)
            i += 1
            continue

        if despida.startswith("---") and set(despida) <= {"-"}:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            borda = OxmlElement("w:pBdr")
            fundo = OxmlElement("w:bottom")
            fundo.set(qn("w:val"), "single")
            fundo.set(qn("w:sz"), "6")
            fundo.set(qn("w:color"), "C9CFD8")
            borda.append(fundo)
            p._p.get_or_add_pPr().append(borda)
            i += 1
            continue

        if despida.startswith("|"):
            bloco = []
            while i < len(linhas) and linhas[i].strip().startswith("|"):
                bloco.append(linhas[i])
                i += 1
            if len(bloco) >= 2:
                _tabela(doc, bloco)
            continue

        if despida.startswith(">"):
            bloco = []
            while i < len(linhas) and linhas[i].strip().startswith(">"):
                bloco.append(linhas[i].strip().lstrip(">").strip())
                i += 1
            _citacao(doc, bloco)
            continue

        if re.match(r"^[-*] ", despida) or re.match(r"^\d+\. ", despida):
            estilo = "List Bullet" if despida[0] in "-*" else "List Number"
            while i < len(linhas) and (
                re.match(r"^[-*] ", linhas[i].strip())
                or re.match(r"^\d+\. ", linhas[i].strip())
            ):
                item = re.sub(r"^([-*]|\d+\.) ", "", linhas[i].strip())
                # Continuação indentada da mesma linha da lista.
                i += 1
                while i < len(linhas) and linhas[i].startswith("  ") and linhas[i].strip():
                    item += " " + linhas[i].strip()
                    i += 1
                p = doc.add_paragraph(style=estilo)
                p.paragraph_format.space_after = Pt(4)
                _escrever_inline(p, item)
            continue

        # Linha de campo — "**Para:** TATICCA" — e paragrafo proprio.
        #
        # Sem isto o bloco de cabecalho da proposta (Para, De, Data, Validade)
        # vira um paragrafo corrido, porque o juntador abaixo une linhas
        # consecutivas. Em markdown elas sao quatro linhas; no Word saiam como
        # uma frase so, e o documento comecava parecendo descuidado.
        if re.match(r"^\*\*[^*]+:\*\*", despida):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _escrever_inline(p, despida)
            i += 1
            continue

        # Parágrafo: junta as linhas até a próxima em branco, porque o markdown
        # quebra em 80 colunas e o Word deve reflurar sozinho.
        bloco = []
        while i < len(linhas) and linhas[i].strip() and not re.match(
            r"^(#|\||>|[-*] |\d+\. |---|\*\*[^*]+:\*\*)", linhas[i].strip()
        ):
            bloco.append(linhas[i].strip())
            i += 1
        _escrever_inline(doc.add_paragraph(), " ".join(bloco))

    destino = origem.with_suffix(".docx")
    doc.save(destino)
    return destino


def main() -> int:
    alvos = [Path(a) for a in sys.argv[1:]] or sorted(PASTA.glob("*.md"))
    if not alvos:
        print("Nenhum .md encontrado em", PASTA)
        return 1
    for origem in alvos:
        destino = converter(origem)
        print(f"  {origem.name}  ->  {destino.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
